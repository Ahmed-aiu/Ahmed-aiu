from docx import Document
import json, os
from datetime import datetime


def generate(rfq_analysis_result,username):
    # Perform Output Handling
    rfq_result = json.loads(rfq_analysis_result)
    rfq_summary = rfq_result['Summary']
    rfq_maincats = rfq_result['Main-Categories']
    rfq_subcats = rfq_result['Sub-Categories']
    rfq_departs = rfq_result['Departments']
    rfq_project_start = rfq_result['Project-Start']
    rfq_project_duration = rfq_result['Project-Duration']
    rfq_contact_tech = rfq_result['Contact-Technical']
    rfq_contact_sales = rfq_result['Contact-Commercial']
    rfq_proposal_deadline = rfq_result['Proposal-Deadline']
    date = datetime.now().date()
    user = username

    # Get the directory of this script
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # Join the script's directory and the file name
    pathTemplate = os.path.join(script_dir, 'RFQReportTemplate_ls240814.docx')
    nameOutput = 'GenAISys_RFQReport_' + datetime.now().strftime("%y%m%d%H%M%S") + '.docx'
    pathOutput = os.path.join(script_dir, nameOutput)

    # Open the .docx file
    with open(pathTemplate, 'rb') as fileTemplate:
        doc = Document(fileTemplate)

        # Define the string patterns and their replacements
        replacements = {
            '<Summary>': rfq_summary,
            '<ContactS>': rfq_contact_sales,
            '<ContactT>': rfq_contact_tech,
            '<PStart>': rfq_project_start,
            '<PDuration>': rfq_project_duration,
            '<Deadline>': rfq_proposal_deadline,
            '<Date>': date,
            '<user-login>': user,
            # Add more patterns and replacements as needed
        }
        # Iterate over the rfq_maincats dictionary
        for i, category in enumerate(rfq_maincats.keys(), start=1):
            # Create a pattern for each activity
            pattern = '<Category' + str(i) + '>'
            # Add the pattern and its replacement to the replacements dictionary
            replacements[pattern] = rfq_maincats[category]
        # Iterate over the rfq_subcats dictionary with enumeration
        for j, category in enumerate(rfq_subcats.keys(), start=1):
            # Create a pattern for each activity
            pattern = '<Activity' + str(j) + '>'
            # Add the pattern and its replacement to the replacements dictionary
            replacements[pattern] = rfq_subcats[category]
        # Iterate over the rfq_departs dictionary
        for i, depart in enumerate(rfq_departs.keys(), start=1):
            # Create a pattern for each key in the dictionary
            for key in rfq_departs[depart].keys():
                pattern = key + str(i)
                # Add the pattern and its replacement to the replacements dictionary
                replacements[pattern] = rfq_departs[depart][key]

        ## Adjust Table Length ##
        # Get the maximum number of relevant rows
        max_rows_work = max(len(rfq_maincats), len(rfq_subcats))
        max_rows_departs = len(rfq_departs)
        # Iterate over all tables in the .docx file and adjust their lenght
        for table in doc.tables:
            # Get the text in the first cell of the row
            first_cell_text = table.rows[0].cells[0].text
            # Identify Content table
            if 'Categories' in first_cell_text:
                # Get the total number of rows in the table
                total_rows = len(table.rows)
                # If the total number of rows is greater than the maximum number of relevant rows
                if total_rows > max_rows_work:
                    # Calculate the number of rows to delete
                    rows_to_delete = total_rows - max_rows_work - 1
                    # Delete the last rows
                    for _ in range(rows_to_delete):
                        row = table._element.getchildren()[-1]
                        table._element.remove(row)
            # Identify Department table
            if 'Department' in first_cell_text:
                # Get the total number of rows in the table
                total_rows = len(table.rows)
                # If the total number of rows is greater than the maximum number of relevant rows
                if total_rows > max_rows_departs:
                    # Calculate the number of rows to delete
                    rows_to_delete = total_rows - max_rows_departs - 1
                    # Delete the last rows
                    for _ in range(rows_to_delete):
                        row = table._element.getchildren()[-1]
                        table._element.remove(row)

        ## replace all remaining placeholders ##
        # Iterate over all paragraphs in the .docx file
        for para in doc.paragraphs:
            # Iterate over all runs in the paragraph
            for run in para.runs:
                # Replace the string patterns with the corresponding strings
                for pattern, replacement in replacements.items():
                    if pattern in run.text:
                        run.text = run.text.replace(pattern, str(replacement))
        # Iterate over all tables in the .docx file
        for table in doc.tables:
            # Iterate over all cells in the table
            for row in table.rows:
                for cell in row.cells:
                    # Iterate over all paragraphs in the cell
                    for para in cell.paragraphs:
                        # Iterate over all runs in the paragraph
                        for run in para.runs:
                            # Replace the string patterns with the corresponding strings
                            for pattern, replacement in replacements.items():
                                if pattern in run.text:
                                    # Replace the pattern in the run text
                                    run.text = run.text.replace(pattern, str(replacement))
                                    # Manually set the formatting for the part of the text that matches the pattern
                                    run._r.getchildren()[0].text = run.text
        
        ## temporary version sending .docx
        # Save the modified .docx file
        doc.save(pathOutput)
        return pathOutput, nameOutput

        ## converting PDF using azure fucntion
        # Convert the .docx file to a .pdf file in Azure Function
        # import requests
        # response = requests.get('https://your-azure-function-url', params={'file_path': file_path})

        # # Get the download URL from the response
        # download_url = response.text

        # # Download the file
        # response = requests.get(download_url)
        # with open('output.pdf', 'wb') as f:
        #     f.write(response.content)

        # # Send the file to the client
        # return send_file('output.pdf', as_attachment=True)